import docx, os

# using python-docx module

# opening a document
document = docx.Document(r"C:\Users\micha\Documents\Training\demo.docx")

# changing the directory
os.chdir(r"C:\Users\micha\Documents\Training")

# printing the text from paragraph 0
print(document.paragraphs[0].text)

p = document.paragraphs[1]

# going through runs
print(p.runs[0].text)

# printing p.runs[1] text
print(p.runs[1].text)

# checking if the text is bolded
print(p.runs[1].bold)

# making the text also italic
p.runs[1].italic = True

# changing the style of the text
p.style = "Title"

# saving the changes
document.save(r"C:\Users\micha\Documents\Training\demo2.docx")

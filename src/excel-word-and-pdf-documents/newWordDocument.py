import docx, os

# Working with a new Word document

# creating a new document
d = docx.Document()

# adding a paragraph
d.add_paragraph("Hello world!")
d.add_paragraph("There is another paragraph.")
d.add_paragraph("And here is another one.")
d.add_paragraph("One more here.")

# adding a run

p = d.paragraphs[0]
p.add_run("This is a new run.")

# printing runs
print(p.runs)

# bolding the second run
p.runs[1].bold = True

# saving the document
d.save(r"C:\Users\micha\Documents\Training\myNewWordDocument.docx")


# defining a function that gets text from a word document
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for paragraph in doc.paragraphs:
        fullText.append(paragraph.text)
    return "\n".join(fullText)


# printing the text from a given Word document
print(getText(r"C:\Users\micha\Documents\Training\myNewWordDocument.docx"))
